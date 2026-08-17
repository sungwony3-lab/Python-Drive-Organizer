import base64
import hashlib
import hmac
import json
import re
import secrets
import sqlite3
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path

from database import PROJECT_ROOT
from search_service import FullTreeResult, SearchService


EXPORT_DIRECTORY = PROJECT_ROOT / "exports"
EXPORT_AUDIT_LOG_PATH = PROJECT_ROOT / "logs" / "tree_export.log"
EXPORT_ID_PATTERN = re.compile(r"^[A-Za-z0-9_-]{32}$")
EXPORT_FORMATS = {"txt", "docx", "xlsx"}
MAX_OPENAI_FILE_BYTES = 10_000_000
MEDIA_TYPES = {
    "txt": "text/plain; charset=utf-8",
    "docx": (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"
    ),
    "xlsx": (
        "application/vnd.openxmlformats-officedocument."
        "spreadsheetml.sheet"
    ),
}


class TreeCursorError(ValueError):
    """Raised when a pagination cursor is invalid or belongs to another tree."""


class TreeCursorStaleError(TreeCursorError):
    """Raised when SQLite changed between pagination calls."""


class TreeExportNotFoundError(FileNotFoundError):
    """Raised when an opaque export ID has no local file."""


class TreeExportFileTooLargeError(ValueError):
    """Raised when an export exceeds the GPT Actions return-file limit."""


@dataclass(frozen=True)
class SnapshotInfo:
    latest_scan_id: str | None
    latest_scan_status: str | None
    scan_finished_at: str | None
    indexed_files: int
    indexed_folders: int


@dataclass(frozen=True)
class TreeExportResult:
    export_id: str
    format: str
    filename: str
    node_count: int
    folder_count: int
    file_count: int
    latest_scan_id: str | None
    latest_scan_status: str | None
    scan_finished_at: str | None
    created_at: str


def utc_text() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def canonical_json(value) -> str:
    return json.dumps(
        value,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    )


def get_snapshot_info(connection: sqlite3.Connection) -> SnapshotInfo:
    scan = connection.execute(
        """
        SELECT scan_id, status, finished_at
        FROM scan_state
        ORDER BY created_at DESC, rowid DESC
        LIMIT 1
        """
    ).fetchone()
    return SnapshotInfo(
        latest_scan_id=scan["scan_id"] if scan else None,
        latest_scan_status=scan["status"] if scan else None,
        scan_finished_at=scan["finished_at"] if scan else None,
        indexed_files=connection.execute(
            "SELECT COUNT(*) FROM files"
        ).fetchone()[0],
        indexed_folders=connection.execute(
            "SELECT COUNT(*) FROM folders"
        ).fetchone()[0],
    )


def build_tree_snapshot(
    connection: sqlite3.Connection,
    *,
    root_folder: str | None,
    include_files: bool,
    max_depth: int | None,
) -> tuple[FullTreeResult, SnapshotInfo]:
    """Read tree nodes and scan metadata within one SQLite read snapshot."""
    started_transaction = not connection.in_transaction
    if started_transaction:
        connection.execute("BEGIN")
    try:
        tree = SearchService(connection).build_full_tree(
            root_folder_id=root_folder,
            max_depth=max_depth,
            include_files=include_files,
        )
        snapshot = get_snapshot_info(connection)
        return tree, snapshot
    finally:
        if started_transaction:
            connection.rollback()


def snapshot_fingerprint(snapshot: SnapshotInfo) -> str:
    return hashlib.sha256(
        canonical_json(asdict(snapshot)).encode("utf-8")
    ).hexdigest()


def tree_options_fingerprint(
    *, root_folder: str | None, include_files: bool, max_depth: int | None
) -> str:
    return hashlib.sha256(
        canonical_json(
            {
                "root_folder": root_folder,
                "include_files": include_files,
                "max_depth": max_depth,
            }
        ).encode("utf-8")
    ).hexdigest()


def _urlsafe_encode(value: bytes) -> str:
    return base64.urlsafe_b64encode(value).decode("ascii").rstrip("=")


def _urlsafe_decode(value: str) -> bytes:
    padding = "=" * (-len(value) % 4)
    return base64.urlsafe_b64decode(value + padding)


def encode_tree_cursor(
    *, offset: int, options_hash: str, snapshot_hash: str, secret: str
) -> str:
    payload = _urlsafe_encode(
        canonical_json(
            {
                "v": 1,
                "offset": offset,
                "options": options_hash,
                "snapshot": snapshot_hash,
            }
        ).encode("utf-8")
    )
    signature = _urlsafe_encode(
        hmac.new(secret.encode("utf-8"), payload.encode("ascii"), hashlib.sha256).digest()
    )
    return f"{payload}.{signature}"


def decode_tree_cursor(
    cursor: str,
    *,
    options_hash: str,
    snapshot_hash: str,
    secret: str,
) -> int:
    try:
        payload, signature = cursor.split(".", 1)
        expected = _urlsafe_encode(
            hmac.new(
                secret.encode("utf-8"), payload.encode("ascii"), hashlib.sha256
            ).digest()
        )
        if not hmac.compare_digest(signature, expected):
            raise TreeCursorError("The tree cursor signature is invalid.")
        data = json.loads(_urlsafe_decode(payload).decode("utf-8"))
        if data.get("v") != 1:
            raise TreeCursorError("The tree cursor version is invalid.")
        offset = data.get("offset")
        if not isinstance(offset, int) or offset < 0:
            raise TreeCursorError("The tree cursor offset is invalid.")
    except TreeCursorError:
        raise
    except (ValueError, TypeError, UnicodeDecodeError, json.JSONDecodeError) as error:
        raise TreeCursorError("The tree cursor is malformed.") from error
    if data.get("options") != options_hash:
        raise TreeCursorError("The tree cursor belongs to different options.")
    if data.get("snapshot") != snapshot_hash:
        raise TreeCursorStaleError("The SQLite tree snapshot changed.")
    return offset


def paginate_tree(
    result: FullTreeResult,
    snapshot: SnapshotInfo,
    *,
    root_folder: str | None,
    include_files: bool,
    max_depth: int | None,
    page_size: int,
    cursor: str | None,
    secret: str,
) -> dict:
    options_hash = tree_options_fingerprint(
        root_folder=root_folder,
        include_files=include_files,
        max_depth=max_depth,
    )
    current_snapshot_hash = snapshot_fingerprint(snapshot)
    offset = 0
    if cursor:
        offset = decode_tree_cursor(
            cursor,
            options_hash=options_hash,
            snapshot_hash=current_snapshot_hash,
            secret=secret,
        )
    if offset > len(result.items):
        raise TreeCursorError("The tree cursor offset exceeds the result.")
    items = result.items[offset : offset + page_size]
    next_offset = offset + len(items)
    has_more = next_offset < len(result.items)
    next_cursor = (
        encode_tree_cursor(
            offset=next_offset,
            options_hash=options_hash,
            snapshot_hash=current_snapshot_hash,
            secret=secret,
        )
        if has_more
        else None
    )
    return {
        "total_nodes": len(result.items),
        "showing": len(items),
        "next_cursor": next_cursor,
        "has_more": has_more,
        "items": items,
        "latest_scan_id": snapshot.latest_scan_id,
        "latest_scan_status": snapshot.latest_scan_status,
        "scan_finished_at": snapshot.scan_finished_at,
    }


def _metadata_lines(snapshot: SnapshotInfo, created_at: str) -> list[str]:
    return [
        "Google Drive Tree",
        f"Generated: {created_at}",
        f"Latest scan ID: {snapshot.latest_scan_id or 'N/A'}",
        f"Latest scan status: {snapshot.latest_scan_status or 'N/A'}",
        f"Scan finished: {snapshot.scan_finished_at or 'N/A'}",
        f"Indexed folders: {snapshot.indexed_folders}",
        f"Indexed files: {snapshot.indexed_files}",
        "Source: SQLite snapshot (not a real-time Google Drive query)",
    ]


def _write_txt(
    path: Path,
    *,
    tree: FullTreeResult,
    snapshot: SnapshotInfo,
    created_at: str,
) -> None:
    content = "\n".join(_metadata_lines(snapshot, created_at))
    content += "\n\n" + tree.text + "\n"
    path.write_text(content, encoding="utf-8")


def _set_docx_run_font(run, name: str, size_pt: float) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = name
    run.font.size = Pt(size_pt)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def _write_docx(
    path: Path,
    *,
    tree: FullTreeResult,
    snapshot: SnapshotInfo,
    created_at: str,
) -> None:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("Google Drive Tree")
    _set_docx_run_font(title_run, "Malgun Gothic", 24)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run(
        "SQLite snapshot export - not a real-time Google Drive query"
    )
    _set_docx_run_font(subtitle_run, "Malgun Gothic", 10)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(89, 89, 89)

    for line in _metadata_lines(snapshot, created_at)[1:7]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(line)
        _set_docx_run_font(run, "Malgun Gothic", 9.5)

    divider = document.add_paragraph()
    divider.paragraph_format.space_before = Pt(8)
    divider.paragraph_format.space_after = Pt(8)
    divider_run = divider.add_run("Tree")
    _set_docx_run_font(divider_run, "Malgun Gothic", 14)
    divider_run.bold = True
    divider_run.font.color.rgb = RGBColor(46, 116, 181)

    for item in tree.items:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(
            min(item["level"], 24) * 0.18
        )
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        label = "[Folder]" if item["node_type"] == "FOLDER" else "[File]"
        run = paragraph.add_run(f"{label} {item['name']}")
        _set_docx_run_font(run, "Malgun Gothic", 8.5)
        if item["node_type"] == "FOLDER":
            run.bold = True
            run.font.color.rgb = RGBColor(31, 78, 121)
        else:
            run.font.color.rgb = RGBColor(64, 64, 64)

    footer = section.footer.paragraphs[0]
    footer.alignment = 2
    footer_run = footer.add_run("Python Drive Organizer | SQLite snapshot")
    _set_docx_run_font(footer_run, "Malgun Gothic", 8)
    footer_run.font.color.rgb = RGBColor(117, 117, 117)
    document.save(path)


def _write_xlsx(
    path: Path,
    *,
    tree: FullTreeResult,
    snapshot: SnapshotInfo,
    created_at: str,
) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Drive Tree"
    sheet.sheet_view.showGridLines = False

    metadata = _metadata_lines(snapshot, created_at)
    for row_number, line in enumerate(metadata, start=1):
        sheet.cell(row=row_number, column=1, value=line)
    sheet.merge_cells("A1:H1")
    sheet["A1"].font = Font(name="Malgun Gothic", size=18, bold=True, color="1F4E79")
    sheet["A1"].alignment = Alignment(vertical="center")
    sheet.row_dimensions[1].height = 28
    for row_number in range(2, 9):
        sheet.merge_cells(start_row=row_number, start_column=1, end_row=row_number, end_column=8)
        sheet.cell(row=row_number, column=1).font = Font(
            name="Malgun Gothic", size=10, color="595959"
        )

    header_row = 10
    headers = [
        "Level",
        "Type",
        "Name",
        "Path",
        "Parent ID",
        "Item ID",
        "MIME Type",
        "Modified Time",
    ]
    for column, header in enumerate(headers, start=1):
        sheet.cell(row=header_row, column=column, value=header)
    header_fill = PatternFill("solid", fgColor="1F4E79")
    thin_gray = Side(style="thin", color="D9E2F3")
    for cell in sheet[header_row]:
        cell.font = Font(name="Malgun Gothic", size=10, bold=True, color="FFFFFF")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = Border(bottom=thin_gray)
    sheet.row_dimensions[header_row].height = 24

    for item in tree.items:
        sheet.append(
            [
                item["level"],
                item["node_type"],
                item["name"],
                item["path"],
                item["parent_id"],
                item["id"],
                item["mime_type"],
                item["modified_time"],
            ]
        )
        row = sheet.max_row
        sheet.row_dimensions[row].outlineLevel = min(item["level"], 7)
        for column in range(1, 9):
            cell = sheet.cell(row=row, column=column)
            cell.font = Font(name="Malgun Gothic", size=9)
            cell.alignment = Alignment(
                horizontal="center" if column in (1, 2) else "left",
                vertical="center",
                wrap_text=column in (3, 4),
            )
        if item["node_type"] == "FOLDER":
            sheet.cell(row=row, column=2).font = Font(
                name="Malgun Gothic", size=9, bold=True, color="1F4E79"
            )

    last_row = max(sheet.max_row, header_row)
    sheet.auto_filter.ref = f"A{header_row}:H{last_row}"
    sheet.freeze_panes = f"A{header_row + 1}"
    widths = {"A": 9, "B": 12, "C": 38, "D": 70, "E": 30, "F": 34, "G": 34, "H": 24}
    for column, width in widths.items():
        sheet.column_dimensions[column].width = width
    workbook.save(path)


def _write_audit(
    log_path: Path,
    *,
    export_id: str,
    export_format: str,
    node_count: int,
    status: str,
) -> None:
    event = {
        "timestamp": utc_text(),
        "export_id": export_id,
        "format": export_format,
        "node_count": node_count,
        "status": status,
    }
    try:
        log_path = Path(log_path)
        log_path.parent.mkdir(parents=True, exist_ok=True)
        with log_path.open("a", encoding="utf-8") as handle:
            handle.write(canonical_json(event) + "\n")
    except OSError:
        return


def create_tree_export(
    connection: sqlite3.Connection,
    *,
    export_format: str,
    root_folder: str | None,
    include_files: bool,
    max_depth: int | None,
    export_directory: Path = EXPORT_DIRECTORY,
    audit_log_path: Path = EXPORT_AUDIT_LOG_PATH,
) -> TreeExportResult:
    normalized_format = export_format.casefold().strip()
    if normalized_format not in EXPORT_FORMATS:
        raise ValueError("format must be txt, docx, or xlsx.")
    tree, snapshot = build_tree_snapshot(
        connection,
        root_folder=root_folder,
        include_files=include_files,
        max_depth=max_depth,
    )
    created_at = utc_text()
    output_directory = Path(export_directory).resolve()
    output_directory.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    while True:
        export_id = secrets.token_urlsafe(24)
        filename = f"drive_tree_{timestamp}_{export_id}.{normalized_format}"
        output_path = output_directory / filename
        if not output_path.exists():
            break
    temporary_path = output_path.with_suffix(output_path.suffix + ".tmp")
    try:
        if normalized_format == "txt":
            _write_txt(
                temporary_path,
                tree=tree,
                snapshot=snapshot,
                created_at=created_at,
            )
        elif normalized_format == "docx":
            _write_docx(
                temporary_path,
                tree=tree,
                snapshot=snapshot,
                created_at=created_at,
            )
        else:
            _write_xlsx(
                temporary_path,
                tree=tree,
                snapshot=snapshot,
                created_at=created_at,
            )
        temporary_path.replace(output_path)
    except Exception:
        temporary_path.unlink(missing_ok=True)
        _write_audit(
            audit_log_path,
            export_id=export_id,
            export_format=normalized_format,
            node_count=len(tree.items),
            status="FAILED",
        )
        raise

    _write_audit(
        audit_log_path,
        export_id=export_id,
        export_format=normalized_format,
        node_count=len(tree.items),
        status="COMPLETED",
    )
    return TreeExportResult(
        export_id=export_id,
        format=normalized_format,
        filename=filename,
        node_count=len(tree.items),
        folder_count=len(tree.folder_ids),
        file_count=len(tree.file_ids),
        latest_scan_id=snapshot.latest_scan_id,
        latest_scan_status=snapshot.latest_scan_status,
        scan_finished_at=snapshot.scan_finished_at,
        created_at=created_at,
    )


def resolve_export_file(
    export_id: str, export_directory: Path = EXPORT_DIRECTORY
) -> tuple[Path, str]:
    if not isinstance(export_id, str) or not EXPORT_ID_PATTERN.fullmatch(export_id):
        raise TreeExportNotFoundError("The export was not found.")
    directory = Path(export_directory).resolve()
    candidates = []
    for extension in EXPORT_FORMATS:
        candidates.extend(directory.glob(f"drive_tree_*_{export_id}.{extension}"))
    candidates = [
        candidate.resolve()
        for candidate in candidates
        if candidate.is_file() and candidate.resolve().parent == directory
    ]
    if len(candidates) != 1:
        raise TreeExportNotFoundError("The export was not found.")
    path = candidates[0]
    return path, MEDIA_TYPES[path.suffix.lstrip(".").casefold()]


def _write_file_return_audit(
    log_path: Path,
    *,
    export_id: str,
    export_format: str,
    byte_size: int,
    status: str,
) -> None:
    event = {
        "timestamp": utc_text(),
        "export_id": export_id,
        "format": export_format,
        "byte_size": byte_size,
        "status": status,
    }
    try:
        log_path = Path(log_path)
        log_path.parent.mkdir(parents=True, exist_ok=True)
        with log_path.open("a", encoding="utf-8") as handle:
            handle.write(canonical_json(event) + "\n")
    except OSError:
        return


def build_openai_file_response(
    export_id: str,
    *,
    export_directory: Path = EXPORT_DIRECTORY,
    audit_log_path: Path = EXPORT_AUDIT_LOG_PATH,
    max_file_bytes: int = MAX_OPENAI_FILE_BYTES,
) -> dict:
    """Return one local export using the GPT Actions openaiFileResponse contract."""
    path, media_type = resolve_export_file(
        export_id, export_directory=export_directory
    )
    export_format = path.suffix.lstrip(".").casefold()
    byte_size = path.stat().st_size
    if byte_size > max_file_bytes:
        _write_file_return_audit(
            audit_log_path,
            export_id=export_id,
            export_format=export_format,
            byte_size=byte_size,
            status="GPT_FILE_TOO_LARGE",
        )
        raise TreeExportFileTooLargeError(
            "The export exceeds the GPT Actions 10 MB file limit."
        )

    file_bytes = path.read_bytes()
    if len(file_bytes) > max_file_bytes:
        _write_file_return_audit(
            audit_log_path,
            export_id=export_id,
            export_format=export_format,
            byte_size=len(file_bytes),
            status="GPT_FILE_TOO_LARGE",
        )
        raise TreeExportFileTooLargeError(
            "The export exceeds the GPT Actions 10 MB file limit."
        )

    response = {
        "openaiFileResponse": [
            {
                "name": path.name,
                "mime_type": media_type.split(";", 1)[0],
                "content": base64.b64encode(file_bytes).decode("ascii"),
            }
        ]
    }
    _write_file_return_audit(
        audit_log_path,
        export_id=export_id,
        export_format=export_format,
        byte_size=len(file_bytes),
        status="OPENAI_FILE_RETURNED",
    )
    return response
